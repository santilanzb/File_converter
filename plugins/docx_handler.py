"""
Concrete implementation of FileHandler for Word documents (.docx).

This module uses python-docx to read and write Word documents.
"""

from docx import Document
from typing import List, Dict, Any

# Try to import from project structure, fall back to dummy classes
try:
    from ..core.exceptions import FileProcessingError
    from .base_handler import FileHandler
except ImportError:
    try:
        # Try absolute imports for installed package
        from core.exceptions import FileProcessingError
        from plugins.base_handler import FileHandler
    except ImportError:
        # For standalone testing, we'll define dummy classes.
        class FileProcessingError(Exception):
            def __init__(self, path, msg):
                super().__init__(f"{msg}: {path}")

        class FileHandler:
            def read(self, file_path: str): pass
            def write(self, file_path: str, data): pass


class DocxHandler(FileHandler):
    """
    Handles reading and writing of Word documents (.docx files).
    
    Extracts text from paragraphs, tables, and other elements.
    """

    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Reads a Word document and extracts text content.

        Args:
            file_path (str): The path to the input Word document.

        Returns:
            List[Dict[str, Any]]: A list where each dictionary represents content
                                 with type and text information.

        Raises:
            FileProcessingError: If the file cannot be found, read, or processed.
        """
        try:
            doc = Document(file_path)
            content_data = []
            
            # Extract paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():  # Only include non-empty paragraphs
                    content_data.append({
                        'type': 'paragraph',
                        'number': para_num,
                        'content': paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row_num, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                content_data.append({
                    'type': 'table',
                    'number': table_num,
                    'content': table_data,
                    'rows': len(table.rows),
                    'columns': len(table.columns)
                })
            
            return content_data
                
        except FileNotFoundError:
            raise FileProcessingError(file_path, "Word document not found")
        except Exception as e:
            raise FileProcessingError(file_path, f"An error occurred while reading the Word document: {e}")

    def write(self, file_path: str, data: List[Dict[str, Any]]) -> None:
        """
        Writes data to a Word document.

        Args:
            file_path (str): The path to the output Word document.
            data (List[Dict[str, Any]]): The data to be written.

        Raises:
            FileProcessingError: If the file cannot be written.
            ValueError: If the data is invalid.
        """
        if not data:
            raise ValueError("Input data for Word document writing cannot be empty.")
        
        try:
            doc = Document()
            
            for item in data:
                if isinstance(item, dict):
                    # Handle different content types
                    if item.get('type') == 'table' and 'content' in item:
                        # Create table
                        table_data = item['content']
                        if table_data and isinstance(table_data, list):
                            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_data in enumerate(row_data):
                                    table.cell(row_idx, col_idx).text = str(cell_data)
                    
                    elif 'content' in item:
                        # Add paragraph
                        text = str(item['content'])
                        if text.strip():
                            para = doc.add_paragraph(text)
                            # Apply style if specified
                            if 'style' in item:
                                try:
                                    para.style = item['style']
                                except:
                                    # If style doesn't exist, keep default
                                    pass
                    
                    elif 'text' in item:
                        text = str(item['text'])
                        if text.strip():
                            doc.add_paragraph(text)
                    
                    else:
                        # Convert entire dict to string
                        text = str(item)
                        if text.strip():
                            doc.add_paragraph(text)
                
                else:
                    # Handle non-dict items
                    text = str(item)
                    if text.strip():
                        doc.add_paragraph(text)
            
            doc.save(file_path)
            
        except Exception as e:
            raise FileProcessingError(file_path, f"An error occurred while writing the Word document: {e}")
